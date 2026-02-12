"""
File Processing Service for LandPPT
Handles document upload and content extraction as specified in requires.md
"""

import os
import re
import logging
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import tempfile

# Document processing libraries
try:
    import docx
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

from ..api.models import FileUploadResponse

logger = logging.getLogger(__name__)


class FileProcessor:
    """Processes uploaded files and extracts content for PPT generation"""
    
    def __init__(self):
        self.supported_formats = {
            '.docx': self._process_docx,
            '.pdf': self._process_pdf,
            '.txt': self._process_txt,
            '.md': self._process_markdown,
            '.jpg': self._process_image,
            '.jpeg': self._process_image,
            '.png': self._process_image,
        }
        
        # Keywords for scenario detection
        self.scenario_keywords = {
            'tourism': ['æ—…æ¸¸', 'æ™¯ç‚¹', 'è¡Œç¨‹', 'æ—…è¡Œ', 'è§‚å…‰', 'åº¦å‡', 'é…’åº—', 'æœºç¥¨', 'å¯¼æ¸¸'],
            'education': ['æ•™è‚²', 'å­¦ä¹ ', 'è¯¾ç¨‹', 'åŸ¹è®­', 'çŸ¥è¯†', 'ç§‘æ™®', 'å„¿ç«¥', 'å­¦ç”Ÿ', 'æ•™å­¦'],
            'analysis': ['åˆ†æ', 'æ•°æ®', 'ç»Ÿè®¡', 'ç ”ç©¶', 'æŠ¥å‘Š', 'è°ƒæŸ¥', 'å›¾è¡¨', 'è¶‹åŠ¿', 'ç»“è®º'],
            'history': ['å†å²', 'å¤ä»£', 'æ–‡åŒ–', 'ä¼ ç»Ÿ', 'é—äº§', 'æ–‡ç‰©', 'æœä»£', 'äº‹ä»¶', 'äººç‰©'],
            'technology': ['æŠ€æœ¯', 'ç§‘æŠ€', 'åˆ›æ–°', 'æ•°å­—', 'æ™ºèƒ½', 'äººå·¥æ™ºèƒ½', 'äº’è”ç½‘', 'è½¯ä»¶', 'ç¡¬ä»¶'],
            'business': ['å•†ä¸š', 'ä¼ä¸š', 'å¸‚åœº', 'è¥é”€', 'é”€å”®', 'ç®¡ç†', 'æˆ˜ç•¥', 'è´¢åŠ¡', 'æŠ•èµ„'],
            'general': ['ä»‹ç»', 'æ¦‚è¿°', 'æ€»ç»“', 'è¯´æ˜', 'å±•ç¤º', 'æ±‡æŠ¥', 'æ¼”ç¤º', 'åˆ†äº«']
        }
    
    async def process_file(
        self,
        file_path: str,
        filename: str,
        *,
        file_processing_mode: Optional[str] = None,
    ) -> FileUploadResponse:
        """Process uploaded file and extract content"""
        try:
            file_ext = Path(filename).suffix.lower()
            file_size = os.path.getsize(file_path)
            
            if file_ext not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {file_ext}")
            
            # Process file based on type
            if file_ext == ".pdf":
                content = await self._process_pdf(file_path, file_processing_mode=file_processing_mode)
            else:
                processor = self.supported_formats[file_ext]
                content = await processor(file_path)
            
            # Extract topics and suggest scenarios
            topics = self._extract_topics(content)
            scenarios = self._suggest_scenarios(content)
            
            return FileUploadResponse(
                filename=filename,
                size=file_size,
                type=file_ext,
                processed_content=content,
                extracted_topics=topics,
                suggested_scenarios=scenarios,
                message=f"æ–‡ä»¶ {filename} å¤„ç†æˆåŠŸï¼Œæå–äº† {len(content)} ä¸ªå­—ç¬¦çš„å†…å®¹"
            )
            
        except Exception as e:
            logger.error(f"Error processing file {filename}: {e}")
            raise ValueError(f"æ–‡ä»¶å¤„ç†å¤±è´¥: {str(e)}")
    
    async def _process_docx(self, file_path: str) -> str:
        """Process DOCX file"""
        if not DOCX_AVAILABLE:
            raise ValueError("DOCX processing not available. Please install python-docx.")

        def _process_docx_sync(file_path: str) -> str:
            """åŒæ­¥å¤„ç†DOCXæ–‡ä»¶ï¼ˆåœ¨çº¿ç¨‹æ± ä¸­è¿è¡Œï¼‰"""
            doc = Document(file_path)
            content_parts = []

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content_parts.append(text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        content_parts.append(" | ".join(row_text))

            return "\n\n".join(content_parts)

        try:
            # åœ¨çº¿ç¨‹æ± ä¸­æ‰§è¡Œæ–‡ä»¶å¤„ç†ä»¥é¿å…é˜»å¡ä¸»æœåŠ¡
            import asyncio
            loop = asyncio.get_running_loop()
            return await loop.run_in_executor(None, _process_docx_sync, file_path)

        except Exception as e:
            logger.error(f"Error processing DOCX file: {e}")
            raise ValueError(f"DOCX æ–‡ä»¶å¤„ç†å¤±è´¥: {str(e)}")
    
    async def _process_pdf(self, file_path: str, *, file_processing_mode: Optional[str] = None) -> str:
        """Process PDF file"""
        # PyPDF2 is only required for the fallback path. In magic_pdf mode we can use MinerU without PyPDF2.
        if not PDF_AVAILABLE and (file_processing_mode or "").lower() != "magic_pdf":
            raise ValueError("PDF processing not available. Please install PyPDF2.")

        def _process_pdf_sync(file_path: str) -> str:
            """åŒæ­¥å¤„ç†PDFæ–‡ä»¶ï¼ˆåœ¨çº¿ç¨‹æ± ä¸­è¿è¡Œï¼‰"""
            content_parts = []

            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)

                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()

                    if text.strip():
                        content_parts.append(text.strip())

            return "\n\n".join(content_parts)

        try:
            # åœ¨çº¿ç¨‹æ± ä¸­æ‰§è¡Œæ–‡ä»¶å¤„ç†ä»¥é¿å…é˜»å¡ä¸»æœåŠ¡
            import asyncio
            loop = asyncio.get_running_loop()

            if (file_processing_mode or "").lower() == "magic_pdf":
                try:
                    from ..core.config import ai_config
                    if getattr(ai_config, "mineru_api_key", None):
                        os.environ["MINERU_API_KEY"] = str(ai_config.mineru_api_key)
                    if getattr(ai_config, "mineru_base_url", None):
                        os.environ["MINERU_BASE_URL"] = str(ai_config.mineru_base_url)
                except Exception:
                    pass

                try:
                    from summeryanyfile.core.magic_pdf_converter import MagicPDFConverter
                    logger.info(f"magic_pdfæ¨¡å¼ï¼šä¼˜å…ˆä½¿ç”¨MinerUå¤„ç†PDF: {Path(file_path).name}")
                    def _mineru_convert_sync() -> str:
                        converter = MagicPDFConverter()
                        md_content, _encoding = converter.convert_pdf_file(file_path)
                        return (md_content or "").strip()

                    md_content = await loop.run_in_executor(None, _mineru_convert_sync)
                    if md_content:
                        return md_content
                    logger.warning("MinerUè¿”å›å†…å®¹ä¸ºç©ºï¼Œå›é€€åˆ°PyPDF2è§£æ")
                except Exception as e:
                    logger.warning(f"MinerUå¤„ç†PDFå¤±è´¥ï¼Œå›é€€åˆ°PyPDF2è§£æ: {e}")

            if not PDF_AVAILABLE:
                raise ValueError("PyPDF2æœªå®‰è£…ï¼Œä¸”MinerUå¤„ç†å¤±è´¥ï¼Œæ— æ³•è§£æPDF")

            return await loop.run_in_executor(None, _process_pdf_sync, file_path)

        except Exception as e:
            logger.error(f"Error processing PDF file: {e}")
            raise ValueError(f"PDF æ–‡ä»¶å¤„ç†å¤±è´¥: {str(e)}")
    
    async def _process_txt(self, file_path: str) -> str:
        """Process TXT file"""
        def _process_txt_sync(file_path: str) -> str:
            """åŒæ­¥å¤„ç†TXTæ–‡ä»¶ï¼ˆåœ¨çº¿ç¨‹æ± ä¸­è¿è¡Œï¼‰"""
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    content = file.read()

                # Try different encodings if UTF-8 fails
                if not content.strip():
                    encodings = ['gbk', 'gb2312', 'latin1']
                    for encoding in encodings:
                        try:
                            with open(file_path, 'r', encoding=encoding) as file:
                                content = file.read()
                            if content.strip():
                                break
                        except:
                            continue

                return content.strip()
            except Exception as e:
                raise e

        try:
            # åœ¨çº¿ç¨‹æ± ä¸­æ‰§è¡Œæ–‡ä»¶å¤„ç†ä»¥é¿å…é˜»å¡ä¸»æœåŠ¡
            import asyncio
            loop = asyncio.get_running_loop()
            return await loop.run_in_executor(None, _process_txt_sync, file_path)

        except Exception as e:
            logger.error(f"Error processing TXT file: {e}")
            raise ValueError(f"TXT æ–‡ä»¶å¤„ç†å¤±è´¥: {str(e)}")
    
    async def _process_markdown(self, file_path: str) -> str:
        """Process Markdown file"""
        def _process_markdown_sync(file_path: str) -> str:
            """åŒæ­¥å¤„ç†Markdownæ–‡ä»¶ï¼ˆåœ¨çº¿ç¨‹æ± ä¸­è¿è¡Œï¼‰"""
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            return content.strip()

        try:
            # åœ¨çº¿ç¨‹æ± ä¸­æ‰§è¡Œæ–‡ä»¶å¤„ç†ä»¥é¿å…é˜»å¡ä¸»æœåŠ¡
            import asyncio
            loop = asyncio.get_running_loop()
            return await loop.run_in_executor(None, _process_markdown_sync, file_path)

        except Exception as e:
            logger.error(f"Error processing Markdown file: {e}")
            raise ValueError(f"Markdown æ–‡ä»¶å¤„ç†å¤±è´¥: {str(e)}")
    
    async def _process_image(self, file_path: str) -> str:
        """Process image file using OCR"""
        if not OCR_AVAILABLE:
            return "å›¾ç‰‡æ–‡ä»¶å·²ä¸Šä¼ ï¼Œä½† OCR åŠŸèƒ½ä¸å¯ç”¨ã€‚è¯·å®‰è£… pytesseract å’Œ PIL ä»¥å¯ç”¨æ–‡å­—è¯†åˆ«ã€‚"

        def _process_image_sync(file_path: str) -> str:
            """åŒæ­¥å¤„ç†å›¾åƒæ–‡ä»¶ï¼ˆåœ¨çº¿ç¨‹æ± ä¸­è¿è¡Œï¼‰"""
            image = Image.open(file_path)

            # Perform OCR
            text = pytesseract.image_to_string(image, lang='chi_sim+eng')

            if not text.strip():
                return "å›¾ç‰‡æ–‡ä»¶å·²å¤„ç†ï¼Œä½†æœªèƒ½è¯†åˆ«å‡ºæ–‡å­—å†…å®¹ã€‚"

            return text.strip()

        try:
            # åœ¨çº¿ç¨‹æ± ä¸­æ‰§è¡Œå›¾åƒå¤„ç†ä»¥é¿å…é˜»å¡ä¸»æœåŠ¡
            import asyncio
            loop = asyncio.get_running_loop()
            return await loop.run_in_executor(None, _process_image_sync, file_path)

        except Exception as e:
            logger.error(f"Error processing image file: {e}")
            return f"å›¾ç‰‡å¤„ç†å¤±è´¥: {str(e)}"
    
    def _extract_topics(self, content: str) -> List[str]:
        """Extract potential topics from content"""
        if not content:
            return []
        
        topics = []
        
        # Extract sentences that might be topics (short, descriptive)
        sentences = re.split(r'[ã€‚ï¼ï¼Ÿ\n]', content)
        
        for sentence in sentences:
            sentence = sentence.strip()
            # Look for topic-like sentences (10-50 characters, no common words)
            if 10 <= len(sentence) <= 50:
                # Avoid sentences with too many common words
                common_words = ['çš„', 'æ˜¯', 'åœ¨', 'æœ‰', 'å’Œ', 'ä¸', 'æˆ–', 'ä½†', 'è€Œ', 'äº†', 'ç€', 'è¿‡']
                common_count = sum(1 for word in common_words if word in sentence)
                
                if common_count <= 2:  # Not too many common words
                    topics.append(sentence)
        
        # Also extract potential titles (lines that are short and at the beginning)
        lines = content.split('\n')
        for i, line in enumerate(lines[:10]):  # Check first 10 lines
            line = line.strip()
            if 5 <= len(line) <= 30 and not line.endswith('ï¼š'):
                topics.append(line)
        
        # Remove duplicates and limit to top 10
        topics = list(dict.fromkeys(topics))[:10]
        
        return topics
    
    def _suggest_scenarios(self, content: str) -> List[str]:
        """Suggest appropriate scenarios based on content"""
        if not content:
            return ['general']
        
        content_lower = content.lower()
        scenario_scores = {}
        
        # Score each scenario based on keyword matches
        for scenario, keywords in self.scenario_keywords.items():
            score = 0
            for keyword in keywords:
                score += content_lower.count(keyword)
            
            if score > 0:
                scenario_scores[scenario] = score
        
        # Sort by score and return top scenarios
        sorted_scenarios = sorted(scenario_scores.items(), key=lambda x: x[1], reverse=True)
        
        # Return top 3 scenarios, or 'general' if no matches
        if sorted_scenarios:
            return [scenario for scenario, score in sorted_scenarios[:3]]
        else:
            return ['general']
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats"""
        return list(self.supported_formats.keys())
    
    def validate_file(self, filename: str, file_size: int, max_size_mb: int = 100) -> Tuple[bool, str]:
        """Validate uploaded file"""
        file_ext = Path(filename).suffix.lower()
        
        # Check file extension
        if file_ext not in self.supported_formats:
            return False, f"ä¸æ”¯æŒçš„æ–‡ä»¶æ ¼å¼: {file_ext}ã€‚æ”¯æŒçš„æ ¼å¼: {', '.join(self.supported_formats.keys())}"
        
        # Check file size
        max_size_bytes = max_size_mb * 1024 * 1024
        if file_size > max_size_bytes:
            return False, f"æ–‡ä»¶å¤§å°è¶…è¿‡é™åˆ¶ ({max_size_mb}MB)ã€‚å½“å‰æ–‡ä»¶å¤§å°: {file_size / 1024 / 1024:.1f}MB"
        
        # Check specific format requirements
        if file_ext == '.docx' and not DOCX_AVAILABLE:
            return False, "DOCX å¤„ç†åŠŸèƒ½ä¸å¯ç”¨ï¼Œè¯·è”ç³»ç®¡ç†å‘˜å®‰è£… python-docx"
        
        if file_ext == '.pdf' and not PDF_AVAILABLE:
            return False, "PDF å¤„ç†åŠŸèƒ½ä¸å¯ç”¨ï¼Œè¯·è”ç³»ç®¡ç†å‘˜å®‰è£… PyPDF2"
        
        if file_ext in ['.jpg', '.jpeg', '.png'] and not OCR_AVAILABLE:
            return True, "å›¾ç‰‡æ–‡ä»¶å¯ä»¥ä¸Šä¼ ï¼Œä½†æ–‡å­—è¯†åˆ«åŠŸèƒ½ä¸å¯ç”¨"
        
        return True, "æ–‡ä»¶éªŒè¯é€šè¿‡"
    
    async def create_ppt_from_content(self, content: str, suggested_topic: str = None) -> Dict[str, Any]:
        """Create PPT generation request from processed content"""
        # Extract or suggest a topic
        if not suggested_topic:
            topics = self._extract_topics(content)
            suggested_topic = topics[0] if topics else "æ–‡æ¡£å†…å®¹å±•ç¤º"
        
        # Suggest scenarios
        scenarios = self._suggest_scenarios(content)
        primary_scenario = scenarios[0] if scenarios else 'general'
        
        # Create a structured outline from content
        sections = self._create_content_sections(content)
        
        return {
            'topic': suggested_topic,
            'scenario': primary_scenario,
            'requirements': f"åŸºäºä¸Šä¼ æ–‡æ¡£å†…å®¹ç”ŸæˆPPTï¼ŒåŒ…å«ä»¥ä¸‹è¦ç‚¹ï¼š\n{content}",
            'uploaded_content': content,
            'suggested_sections': sections,
            'language': 'zh'
        }
    
    def _create_content_sections(self, content: str) -> List[Dict[str, str]]:
        """Create structured sections from content"""
        sections = []

        # Split content into logical sections
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]

        # Create title slide
        sections.append({
            'type': 'title',
            'title': 'æ–‡æ¡£å†…å®¹å±•ç¤º',
            'subtitle': 'åŸºäºä¸Šä¼ æ–‡æ¡£ç”Ÿæˆ'
        })

        # Create content slides (max 10)
        for i, paragraph in enumerate(paragraphs[:9]):
            if len(paragraph) > 50:  # Only use substantial paragraphs
                # Try to extract a title from the first sentence
                sentences = paragraph.split('ã€‚')
                title = sentences[0][:30] + '...' if len(sentences[0]) > 30 else sentences[0]

                sections.append({
                    'type': 'content',
                    'title': title or f'å†…å®¹ {i+1}',
                    'content': paragraph[:300] + '...' if len(paragraph) > 300 else paragraph
                })

        # Add thank you slide
        sections.append({
            'type': 'thankyou',
            'title': 'è°¢è°¢è§‚çœ‹',
            'subtitle': 'åŸºäºæ–‡æ¡£å†…å®¹ç”Ÿæˆ'
        })

        return sections

    def merge_multiple_files_to_markdown(self, files_content: List[Dict[str, str]]) -> str:
        """
        å°†å¤šä¸ªæ–‡ä»¶çš„å†…å®¹åˆå¹¶ä¸ºä¸€ä¸ªå®Œæ•´çš„Markdownæ–‡æ¡£

        Args:
            files_content: æ–‡ä»¶å†…å®¹åˆ—è¡¨ï¼Œæ¯é¡¹åŒ…å« filename å’Œ content

        Returns:
            åˆå¹¶åçš„Markdownæ ¼å¼å†…å®¹
        """
        if not files_content:
            return ""

        # å¦‚æœåªæœ‰ä¸€ä¸ªæ–‡ä»¶ï¼Œç›´æ¥è¿”å›å…¶å†…å®¹
        if len(files_content) == 1:
            return files_content[0]["content"]

        # æ„å»ºåˆå¹¶åçš„Markdownæ–‡æ¡£
        merged_parts = []

        # æ·»åŠ æ–‡æ¡£æ ‡é¢˜
        merged_parts.append("# åˆå¹¶æ–‡æ¡£å†…å®¹\n")
        merged_parts.append(f"*æœ¬æ–‡æ¡£ç”± {len(files_content)} ä¸ªæºæ–‡ä»¶åˆå¹¶ç”Ÿæˆ*\n")
        merged_parts.append("---\n")

        # æ·»åŠ ç›®å½•
        merged_parts.append("## ğŸ“‹ æ–‡æ¡£ç›®å½•\n")
        for i, file_info in enumerate(files_content, 1):
            filename = file_info["filename"]
            merged_parts.append(f"{i}. [{filename}](#{self._sanitize_anchor(filename)})\n")
        merged_parts.append("\n---\n")

        # æ·»åŠ æ¯ä¸ªæ–‡ä»¶çš„å†…å®¹
        for i, file_info in enumerate(files_content, 1):
            filename = file_info["filename"]
            content = file_info["content"]

            # æ·»åŠ æ–‡ä»¶æ ‡é¢˜ï¼ˆä½œä¸ºä¸€çº§æ ‡é¢˜ï¼‰
            merged_parts.append(f"\n## {i}. {filename} {{#{self._sanitize_anchor(filename)}}}\n")

            # æ·»åŠ åˆ†éš”çº¿
            merged_parts.append("---\n")

            # æ·»åŠ æ–‡ä»¶å†…å®¹
            # å¦‚æœå†…å®¹å·²ç»åŒ…å«Markdownæ ¼å¼ï¼Œä¿æŒåŸæ ·
            # å¦åˆ™å°†å…¶æ ¼å¼åŒ–ä¸ºæ®µè½
            if content.strip():
                # æ£€æŸ¥æ˜¯å¦å·²ç»æ˜¯Markdownæ ¼å¼
                if self._is_markdown_formatted(content):
                    merged_parts.append(f"{content}\n")
                else:
                    # å°†çº¯æ–‡æœ¬å†…å®¹è½¬æ¢ä¸ºæ®µè½
                    paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
                    for paragraph in paragraphs:
                        merged_parts.append(f"{paragraph}\n\n")
            else:
                merged_parts.append("*ï¼ˆæ­¤æ–‡ä»¶æ— å†…å®¹æˆ–å†…å®¹æå–å¤±è´¥ï¼‰*\n")

            # æ·»åŠ æ–‡ä»¶ç»“æŸæ ‡è®°
            merged_parts.append("\n")

        # æ·»åŠ æ–‡æ¡£ç»“å°¾
        merged_parts.append("\n---\n")
        merged_parts.append("*æ–‡æ¡£ç»“æŸ*\n")

        return "".join(merged_parts)

    def _sanitize_anchor(self, text: str) -> str:
        """å°†æ–‡æœ¬è½¬æ¢ä¸ºåˆæ³•çš„Markdowné”šç‚¹"""
        # ç§»é™¤ç‰¹æ®Šå­—ç¬¦ï¼Œåªä¿ç•™å­—æ¯æ•°å­—å’Œä¸­æ–‡
        sanitized = re.sub(r'[^\w\u4e00-\u9fff-]', '-', text)
        # ç§»é™¤å¤šä½™çš„è¿å­—ç¬¦
        sanitized = re.sub(r'-+', '-', sanitized)
        # ç§»é™¤é¦–å°¾çš„è¿å­—ç¬¦
        sanitized = sanitized.strip('-')
        return sanitized.lower()

    def _is_markdown_formatted(self, content: str) -> bool:
        """æ£€æŸ¥å†…å®¹æ˜¯å¦å·²ç»æ˜¯Markdownæ ¼å¼"""
        # ç®€å•æ£€æŸ¥æ˜¯å¦åŒ…å«å¸¸è§çš„Markdownè¯­æ³•
        markdown_indicators = [
            r'^#{1,6}\s',  # æ ‡é¢˜
            r'\*\*.*\*\*',  # ç²—ä½“
            r'\*.*\*',  # æ–œä½“
            r'^\s*[-*+]\s',  # åˆ—è¡¨
            r'^\s*\d+\.\s',  # æœ‰åºåˆ—è¡¨
            r'\[.*\]\(.*\)',  # é“¾æ¥
            r'```',  # ä»£ç å—
        ]

        for indicator in markdown_indicators:
            if re.search(indicator, content, re.MULTILINE):
                return True

        return False
