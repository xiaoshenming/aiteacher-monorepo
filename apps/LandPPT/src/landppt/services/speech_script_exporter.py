"""
Speech Script Export Service
Handles exporting speech scripts to various document formats (DOCX, Markdown)
"""

import logging
import tempfile
import os
from typing import List, Dict, Any, Optional
from pathlib import Path
import asyncio
from concurrent.futures import ThreadPoolExecutor

from .speech_script_service import SlideScriptData

logger = logging.getLogger(__name__)

# Check for optional dependencies
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available. DOCX export will be disabled.")


class SpeechScriptExporter:
    """Service for exporting speech scripts to various document formats"""
    
    def __init__(self):
        self.executor = ThreadPoolExecutor(max_workers=2)
    
    async def export_to_docx(
        self,
        scripts: List[SlideScriptData],
        project_title: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> bytes:
        """Export speech scripts to DOCX format"""
        if not DOCX_AVAILABLE:
            raise ValueError("DOCX export not available. Please install python-docx: pip install python-docx")
        
        # Run the synchronous DOCX generation in thread pool
        loop = asyncio.get_event_loop()
        docx_content = await loop.run_in_executor(
            self.executor,
            self._generate_docx_sync,
            scripts,
            project_title,
            metadata
        )
        
        return docx_content
    
    async def export_to_markdown(
        self,
        scripts: List[SlideScriptData],
        project_title: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> str:
        """Export speech scripts to Markdown format"""
        
        # Run the markdown generation in thread pool
        loop = asyncio.get_event_loop()
        markdown_content = await loop.run_in_executor(
            self.executor,
            self._generate_markdown_sync,
            scripts,
            project_title,
            metadata
        )
        
        return markdown_content
    
    def _generate_docx_sync(
        self,
        scripts: List[SlideScriptData],
        project_title: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> bytes:
        """Generate DOCX document synchronously"""
        try:
            # Create new document
            doc = Document()

            # Set document language and encoding
            doc.core_properties.language = 'zh-CN'

            # Add title
            title = doc.add_heading(f'{project_title} - 演讲稿', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Set font for the title to support Chinese characters
            for run in title.runs:
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', 'Microsoft YaHei')
            
            # Add metadata if provided
            if metadata:
                doc.add_paragraph()
                meta_para = doc.add_paragraph()
                meta_run = meta_para.add_run('生成信息：')
                meta_run.bold = True
                self._set_chinese_font(meta_run)

                if 'generation_time' in metadata:
                    import time
                    gen_time = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(metadata['generation_time']))
                    time_para = doc.add_paragraph(f'生成时间：{gen_time}')
                    self._set_paragraph_chinese_font(time_para)

                if 'total_estimated_duration' in metadata:
                    duration_para = doc.add_paragraph(f'预计总时长：{metadata["total_estimated_duration"]}')
                    self._set_paragraph_chinese_font(duration_para)

                if 'customization' in metadata:
                    custom = metadata['customization']
                    if 'tone' in custom:
                        tone_para = doc.add_paragraph(f'语调风格：{custom["tone"]}')
                        self._set_paragraph_chinese_font(tone_para)
                    if 'target_audience' in custom:
                        audience_para = doc.add_paragraph(f'目标受众：{custom["target_audience"]}')
                        self._set_paragraph_chinese_font(audience_para)

            doc.add_page_break()
            
            # Add scripts for each slide
            for i, script in enumerate(scripts):
                # Add slide header
                if script.slide_index == -1:
                    # Opening remarks
                    slide_header = doc.add_heading('开场白', level=1)
                elif script.slide_index >= len(scripts) - 1 and script.slide_title == "结束语":
                    # Closing remarks
                    slide_header = doc.add_heading('结束语', level=1)
                else:
                    # Regular slide
                    slide_header = doc.add_heading(f'第{script.slide_index + 1}页：{script.slide_title}', level=1)

                # Set Chinese font for header
                self._set_paragraph_chinese_font(slide_header)

                # Add duration if available
                if script.estimated_duration:
                    duration_para = doc.add_paragraph()
                    duration_run = duration_para.add_run(f'预计时长：{script.estimated_duration}')
                    duration_run.italic = True
                    self._set_chinese_font(duration_run)

                # Add script content
                doc.add_paragraph()
                script_para = doc.add_paragraph(script.script_content)
                script_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                self._set_paragraph_chinese_font(script_para)

                # Add speaker notes if available
                if script.speaker_notes:
                    doc.add_paragraph()
                    notes_para = doc.add_paragraph()
                    notes_title_run = notes_para.add_run('演讲提示：')
                    notes_title_run.bold = True
                    self._set_chinese_font(notes_title_run)

                    notes_content_run = notes_para.add_run(script.speaker_notes)
                    notes_content_run.italic = True
                    self._set_chinese_font(notes_content_run)

                # Add page break except for the last script
                if i < len(scripts) - 1:
                    doc.add_page_break()
            
            # Save to temporary file and read content
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                doc.save(temp_file.name)
                temp_file_path = temp_file.name
            
            try:
                with open(temp_file_path, 'rb') as f:
                    content = f.read()
                return content
            finally:
                # Clean up temporary file
                try:
                    os.unlink(temp_file_path)
                except:
                    pass
                    
        except Exception as e:
            logger.error(f"Error generating DOCX: {e}")
            raise

    def _set_chinese_font(self, run):
        """Set Chinese font for a run to prevent encoding issues"""
        try:
            run.font.name = 'Microsoft YaHei'
            # Set East Asian font for Chinese characters
            run._element.rPr.rFonts.set(
                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia',
                'Microsoft YaHei'
            )
        except Exception as e:
            logger.warning(f"Failed to set Chinese font: {e}")

    def _set_paragraph_chinese_font(self, paragraph):
        """Set Chinese font for all runs in a paragraph"""
        try:
            for run in paragraph.runs:
                self._set_chinese_font(run)
        except Exception as e:
            logger.warning(f"Failed to set paragraph Chinese font: {e}")
    
    def _generate_markdown_sync(
        self,
        scripts: List[SlideScriptData],
        project_title: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> str:
        """Generate Markdown document synchronously"""
        try:
            lines = []
            
            # Add title
            lines.append(f'# {project_title} - 演讲稿\n')
            
            # Add metadata if provided
            if metadata:
                lines.append('## 生成信息\n')
                
                if 'generation_time' in metadata:
                    import time
                    gen_time = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(metadata['generation_time']))
                    lines.append(f'- **生成时间**：{gen_time}')
                
                if 'total_estimated_duration' in metadata:
                    lines.append(f'- **预计总时长**：{metadata["total_estimated_duration"]}')
                
                if 'customization' in metadata:
                    custom = metadata['customization']
                    if 'tone' in custom:
                        lines.append(f'- **语调风格**：{custom["tone"]}')
                    if 'target_audience' in custom:
                        lines.append(f'- **目标受众**：{custom["target_audience"]}')
                
                lines.append('')
            
            lines.append('---\n')
            
            # Add scripts for each slide
            for script in scripts:
                # Add slide header
                if script.slide_index == -1:
                    # Opening remarks
                    lines.append('## 开场白\n')
                elif script.slide_index >= len(scripts) - 1 and script.slide_title == "结束语":
                    # Closing remarks
                    lines.append('## 结束语\n')
                else:
                    # Regular slide
                    lines.append(f'## 第{script.slide_index + 1}页：{script.slide_title}\n')
                
                # Add duration if available
                if script.estimated_duration:
                    lines.append(f'**预计时长**：{script.estimated_duration}\n')
                
                # Add script content
                lines.append(script.script_content)
                lines.append('')
                
                # Add speaker notes if available
                if script.speaker_notes:
                    lines.append('> **演讲提示**：' + script.speaker_notes)
                    lines.append('')
                
                lines.append('---\n')
            
            return '\n'.join(lines)
            
        except Exception as e:
            logger.error(f"Error generating Markdown: {e}")
            raise
    
    def is_docx_available(self) -> bool:
        """Check if DOCX export is available"""
        return DOCX_AVAILABLE
    
    async def cleanup(self):
        """Cleanup resources"""
        if hasattr(self, 'executor'):
            self.executor.shutdown(wait=True)


# Global instance
_speech_script_exporter = None

def get_speech_script_exporter() -> SpeechScriptExporter:
    """Get global speech script exporter instance"""
    global _speech_script_exporter
    if _speech_script_exporter is None:
        _speech_script_exporter = SpeechScriptExporter()
    return _speech_script_exporter
